import logging
from pathlib import Path

from docx import Document

from app.parsers.base import BaseParser, ParsedDocument

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    async def parse(self, source_config: dict) -> list[ParsedDocument]:
        file_path = source_config.get("path", "")
        if not Path(file_path).exists():
            logger.error("DOCX file not found: %s", file_path)
            return []

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        if not paragraphs:
            return []

        return [ParsedDocument(
            text="\n\n".join(paragraphs),
            metadata={
                "source": "docx",
                "file": Path(file_path).name,
                "title": source_config.get("name", Path(file_path).stem),
            },
            source_type="docx",
        )]
